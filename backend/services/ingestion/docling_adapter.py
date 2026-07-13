"""
Docling adapter — thin httpx client for the docling sidecar.

The sidecar (`docling_svc/`) wraps IBM Docling so the backend image stays
free of torch / transformers / accelerate. Backend code only ever sees the
adapter's `parse_document(...) -> DoclingParseResult` contract.

Responsibilities split:
  • Sidecar parses bytes → DoclingDocument → flat sections + markdown.
  • PDF uploads are text-first. Local pypdf extraction is the only PDF text
    recovery path; OCR is disabled by policy and `do_ocr=True` is ignored.
  • Adapter does ONE pre-processing step: when the upload looks like a
    structurally-implicit `.txt` file, run `inject_synthetic_headers` first
    so docling sees real `#`/`##` markers and can promote tier_b_plus.
  • Adapter then asks the sidecar to parse the (possibly augmented) bytes
    and packages the response into a `DoclingParseResult` ready for
    `source_classifier` and `tier_chunker` to consume.
"""

from __future__ import annotations

import logging
import os
import re
import csv
import tempfile
from io import BytesIO, StringIO
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path

import httpx

from models.schemas import SourceTier
from services.ingestion.b_plus_normalizer import inject_synthetic_headers
from services.ingestion.bibliographic import (
    FRONTMATTER_DATE_KEYS,
    FRONTMATTER_LANGUAGE_KEYS,
    KIND_FILE_CREATION,
    KIND_PUBLICATION,
    KIND_REVISION,
    DateCandidate,
    build_provenance,
    normalize_language,
    resolve_document_dates,
)

logger = logging.getLogger(__name__)

DOCLING_URL = os.getenv("DOCLING_URL", "http://docling:8500")
# Sidecar timeout. OCR is disabled; this mainly protects large layout parses
# for non-PDF formats and unusual PDFs routed to Docling without OCR.
DOCLING_TIMEOUT_SECONDS = float(os.getenv("DOCLING_TIMEOUT_SECONDS", "600"))
DOCLING_SIDECAR_POLICY = os.getenv("DOCLING_SIDECAR_POLICY", "auto").strip().lower()
DOCLING_AUTO_UNLOAD_AFTER_PARSE = (
    os.getenv("DOCLING_AUTO_UNLOAD_AFTER_PARSE", "false").strip().lower()
    in {"1", "true", "yes", "on"}
)

_PLAIN_TEXT_MIMES = {"text/plain"}
_PLAIN_TEXT_EXTS = {".txt", ".text", ".log"}
_MARKDOWN_MIMES = {"text/markdown", "text/x-markdown"}
_MARKDOWN_EXTS = {".md", ".markdown", ".mdown", ".mkd"}
_HTML_MIMES = {"text/html", "application/xhtml+xml"}
_HTML_EXTS = {".html", ".htm", ".xhtml"}
_EPUB_MIMES = {"application/epub+zip", "application/epub"}
_EPUB_EXTS = {".epub"}
_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_CSV_MIMES = {"text/csv", "application/csv", "text/tab-separated-values"}
_CSV_EXTS = {".csv", ".tsv"}
_SUBTITLE_EXTS = {".vtt", ".srt"}  # router 4 — subtitle/caption transcripts
_SPREADSHEET_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "application/vnd.ms-excel.sheet.macroenabled.12",
}
_SPREADSHEET_EXTS = {".xlsx", ".xlsm", ".xltx", ".xltm", ".xls"}
_BINARY_DOC_EXTS = {
    ".doc", ".docx", ".ppt", ".pptx", ".xls", ".xlsx", ".xlsm", ".epub", ".odt",
    ".ods", ".odp", ".rtf", ".msg", ".eml",
}
# Code lane Phase 1 — extension → tree-sitter language tag. Files matching
# any of these extensions bypass the Docling sidecar entirely via the early-
# intercept gate at the top of parse_document(). The chunker then routes
# them through code_splitter.pack() which respects EMBEDDER_SAFE_MAX_TOKENS.
# Detection is O(1) (single suffix lookup); doc_id is sha256 of raw text
# downstream, so idempotency is preserved across runs.
_CODE_EXT_TO_LANGUAGE: dict[str, str] = {
    # ─── Mainstream programming languages ───────────────────────────────
    ".py": "python",   ".pyi": "python",
    ".rs": "rust",
    ".js": "javascript", ".mjs": "javascript", ".cjs": "javascript", ".jsx": "javascript",
    ".ts": "typescript", ".tsx": "tsx",
    ".go": "go",
    ".lua": "lua",     ".luau": "luau",
    ".cpp": "cpp",     ".cc": "cpp",       ".cxx": "cpp",      ".hpp": "cpp",
    ".c": "c",         ".h": "c",
    ".cu": "cuda",     ".cuh": "cuda",                          # CUDA kernels
    ".java": "java",
    ".rb": "ruby",
    ".php": "php",
    ".swift": "swift",
    ".kt": "kotlin",   ".kts": "kotlin",
    ".scala": "scala",
    ".sh": "bash",     ".bash": "bash",    ".zsh": "bash",
    ".ps1": "powershell",
    ".sql": "sql",
    ".r": "r",
    ".cs": "csharp",
    ".ex": "elixir",   ".exs": "elixir",
    ".hs": "haskell",
    ".ml": "ocaml",
    ".cljs": "clojure",".clj": "clojure",  ".edn": "clojure",
    ".dart": "dart",
    ".zig": "zig",
    ".nix": "nix",
    ".m": "objc",      ".mm": "objc",                           # Objective-C / Objective-C++
    # ─── Shaders ────────────────────────────────────────────────────────
    ".glsl": "glsl",   ".frag": "glsl",    ".vert": "glsl",
    ".hlsl": "hlsl",
    # ─── Web frameworks ─────────────────────────────────────────────────
    ".vue": "vue",
    ".svelte": "svelte",
    # ─── Styling ────────────────────────────────────────────────────────
    # HTML uploads default to the local_html prose extractor below. That is
    # the safer RAG default for web/document exports: strip navigation,
    # scripts, style, and markup before chunking. Source-code HTML should be
    # wrapped in a code fence inside markdown or renamed before ingest.
    ".css": "css",     ".scss": "css",
    # ─── XML family ─────────────────────────────────────────────────────
    # Apple property lists, Storyboards, Xibs, entitlements; Roblox place
    # & model files (decomposed XML format from rbx-dom).
    ".xml": "xml",
    ".plist": "xml",   ".storyboard": "xml",
    ".xib": "xml",     ".entitlements": "xml",
    ".rbxmx": "xml",   ".rbxlx": "xml",
    # ─── Data / config ──────────────────────────────────────────────────
    # Parsed by the pack so we get atomic packing (never mid-record splits)
    # and Ghost B skips them. symbols_defined will usually be empty —
    # these are data, not callable code.
    ".json": "json",   ".jsonl": "json",
    ".ipynb": "json",  # Jupyter notebooks (cell-level extraction is Phase 2)
    ".yaml": "yaml",   ".yml": "yaml",
    ".toml": "toml",   # pyproject.toml, Cargo.toml, etc.
    ".ini": "ini",     ".cfg": "ini",
    # ─── IaC / build tooling ────────────────────────────────────────────
    ".tf": "hcl",      ".tfvars": "hcl",   ".hcl": "hcl",
    # ─── API schemas ────────────────────────────────────────────────────
    ".proto": "proto",                                          # Protocol Buffers
    ".graphql": "graphql", ".gql": "graphql",
}

# Filenames with no extension that should route through the code lane.
# Looked up case-insensitively before the extension map fires.
_CODE_FILENAME_TO_LANGUAGE: dict[str, str] = {
    "dockerfile":     "dockerfile",
    "containerfile":  "dockerfile",  # Podman / OCI alias
    "makefile":       "make",
    "gnumakefile":    "make",
    "cmakelists.txt": "cmake",
}
_FAST_PDF_MIN_TOTAL_CHARS = 1200
_FAST_PDF_MIN_AVG_CHARS_PER_PAGE = 80
_FAST_PDF_MIN_NONEMPTY_PAGE_RATIO = 0.25
_FAST_PDF_MAX_REPLACEMENT_RATIO = 0.03
_TABLE_PARSE_MAX_SHEETS = int(os.getenv("TABLE_PARSE_MAX_SHEETS", "20"))
_TABLE_PARSE_MAX_ROWS_PER_SHEET = int(os.getenv("TABLE_PARSE_MAX_ROWS_PER_SHEET", "5000"))


def _looks_like_pdf(filename: str, mime: str) -> bool:
    return (mime or "").lower() == "application/pdf" or (filename or "").lower().endswith(".pdf")


def _extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def _looks_like_markdown(filename: str, mime: str) -> bool:
    return (mime or "").lower() in _MARKDOWN_MIMES or _extension(filename) in _MARKDOWN_EXTS


# YAML frontmatter: `---` on line 1, closing `---`/`...` within the first 4 KB.
# Scraped/exported markdown (e.g. the merged corpus: source_url / site_name /
# priority / extracted headers) carries it on every file. Left in place it
# leaks into the first chunk as embedding noise AND fires fact cues — the
# `extracted: 2026-03-24` line became a confident-but-junk timestamp fact in
# the Phase A smoke. It is metadata, not content; strip before sectioning.
_FRONTMATTER_RE = re.compile(
    r"\A﻿?---[ \t]*\r?\n.{0,4096}?\r?\n(?:---|\.\.\.)[ \t]*\r?\n",
    re.DOTALL,
)


_HTML_TITLE_RE = re.compile(r"<title[^>]*>(.*?)</title>", re.I | re.S)

_HTML_PUBLISHED_KEYS = {
    "article:published_time",
    "citation_publication_date",
    "citation_date",
    "datepublished",
    "publish_date",
    "publish-date",
    "publishdate",
}


class _HTMLMetadataParser(HTMLParser):
    """Order-independent extraction of explicit HTML metadata attributes."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.author: str | None = None
        self.language: str | None = None
        self.published: tuple[str, str] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        values = {str(k).lower(): v for k, v in attrs if k}
        tag = tag.lower()
        if tag == "html" and self.language is None:
            self.language = values.get("lang")
            return
        if tag != "meta":
            return
        key = str(values.get("property") or values.get("name") or "").lower()
        content = str(values.get("content") or "").strip()
        if key == "author" and content and self.author is None:
            self.author = content
        if key in _HTML_PUBLISHED_KEYS and content and self.published is None:
            self.published = (key, content)


def _meta_from_frontmatter(text: str) -> dict:
    """M2 — read title/author/date out of the YAML frontmatter BEFORE it is
    stripped (it was provenance being thrown away). Line-level parse, no yaml
    dependency; quoted values unwrapped."""
    clean = text.lstrip("\ufeff")
    if not clean.startswith("---"):
        return {}
    m = _FRONTMATTER_RE.match(clean)
    if not m:
        return {}
    out: dict = {}
    candidates: list[DateCandidate] = []
    seen_date_keys: set[str] = set()
    for line in m.group(0).splitlines()[1:-1]:
        k, _, v = line.partition(":")
        k, v = k.strip().lower(), v.strip().strip("'\"")
        if not v or ":" not in line:
            continue
        if k == "title" and "title" not in out:
            out["title"] = v[:300]
        elif k in ("author", "authors", "creator", "by") and "author" not in out:
            out["author"] = v[:200]
        elif k in FRONTMATTER_LANGUAGE_KEYS and "language_meta" not in out:
            lang = normalize_language(v)
            if lang:
                out["language_meta"] = lang
        elif k in FRONTMATTER_DATE_KEYS and k not in seen_date_keys:
            # T-HOOK-3: no more conflation into document_date — each dated key
            # becomes a typed candidate (published→publication, date→ambiguous,
            # created/modified→file time); finalize_source_meta resolves.
            seen_date_keys.add(k)
            kind, method = FRONTMATTER_DATE_KEYS[k]
            candidates.append(
                DateCandidate(raw=v[:80], kind=kind, method=method,
                              source=f"frontmatter:{k}")
            )
    if candidates:
        out["date_candidates"] = candidates
    return out


def _meta_from_html(raw_bytes: bytes) -> dict:
    try:
        head = raw_bytes[:65536].decode("utf-8", errors="replace")
    except Exception:
        return {}
    out: dict = {}
    m = _HTML_TITLE_RE.search(head)
    if m:
        t = re.sub(r"\s+", " ", m.group(1)).strip()
        if t:
            out["title"] = t[:300]
    parser = _HTMLMetadataParser()
    try:
        parser.feed(head)
    except Exception:
        # Malformed scraped HTML must not fail ingestion; title regex above and
        # all downstream fallbacks remain available.
        pass
    if parser.author:
        out["author"] = parser.author[:200]
    lang = normalize_language(parser.language)
    if lang:
        out["language_meta"] = lang
    if parser.published:
        key, raw_date = parser.published
        out["date_candidates"] = [
            DateCandidate(
                raw=raw_date[:80],
                kind=KIND_PUBLICATION,
                method="html_meta_published",
                source=f"html_meta:{key}",
            )
        ]
    return out


def _meta_from_docx(raw_bytes: bytes) -> dict:
    try:
        from docx import Document

        cp = Document(BytesIO(raw_bytes)).core_properties
        out: dict = {}
        title = getattr(cp, "title", None)
        author = getattr(cp, "author", None)
        if title:
            out["title"] = str(title).strip()[:300]
        if author:
            out["author"] = str(author).strip()[:200]
        lang = normalize_language(getattr(cp, "language", None))
        if lang:
            out["language_meta"] = lang
        # T-HOOK-3 de-conflation: core-props created/modified are FILE times,
        # never publication — they may only ever explain a null date.
        candidates: list[DateCandidate] = []
        created = getattr(cp, "created", None)
        if created:
            candidates.append(
                DateCandidate(
                    raw=created.date().isoformat(),
                    kind=KIND_FILE_CREATION,
                    method="docx_core_created",
                    source="docx:core_properties.created",
                )
            )
        if getattr(cp, "modified", None):
            candidates.append(
                DateCandidate(
                    raw=cp.modified.date().isoformat(),
                    kind=KIND_REVISION,
                    method="docx_core_modified",
                    source="docx:core_properties.modified",
                )
            )
        if candidates:
            out["date_candidates"] = candidates
        return out
    except Exception:
        return {}


def _meta_from_pdf(raw_bytes: bytes) -> dict:
    try:
        from pypdf import PdfReader

        md = PdfReader(BytesIO(raw_bytes)).metadata or {}
        out: dict = {}
        t = str(md.get("/Title") or "").strip()
        a = str(md.get("/Author") or "").strip()
        if t:
            out["title"] = t[:300]
        if a:
            out["author"] = a[:200]
        # T-HOOK-3 de-conflation: /CreationDate and /ModDate are FILE times,
        # never publication — they may only ever explain a null date.
        candidates: list[DateCandidate] = []
        for key, kind, method in (
            ("/CreationDate", KIND_FILE_CREATION, "pdf_creation_date"),
            ("/ModDate", KIND_REVISION, "pdf_mod_date"),
        ):
            d = str(md.get(key) or "").strip()
            if d.startswith("D:") and len(d) >= 10 and d[2:10].isdigit():
                candidates.append(
                    DateCandidate(
                        raw=f"{d[2:6]}-{d[6:8]}-{d[8:10]}",
                        kind=kind,
                        method=method,
                        source=f"pdf:{key}",
                    )
                )
        if candidates:
            out["date_candidates"] = candidates
        return out
    except Exception:
        return {}


def _apply_meta(result: "DoclingParseResult", meta: dict) -> None:
    for k in ("title", "author", "language_meta"):
        if meta.get(k) and not getattr(result, k, None):
            setattr(result, k, meta[k])
    if meta.get("date_candidates"):
        current = getattr(result, "date_candidates", None)
        if current is None:
            setattr(result, "date_candidates", list(meta["date_candidates"]))
        else:
            current.extend(meta["date_candidates"])
        # Preserve the parse-result convenience contract, but resolve through
        # the de-conflated candidate model: publication metadata is visible
        # immediately while file-creation/revision times remain null.
        result.document_date = resolve_document_dates(
            getattr(result, "date_candidates", None) or []
        )["document_date"]


_SOURCE_TYPE_BY_FORMAT = {
    "local_html": "webpage",
    "pypdf_fast_text": "pdf",
    "local_docx": "document",
    "local_markdown": "markdown",
    "local_text": "text",
    "youtube_transcript": "transcript",
    "subtitle_vtt": "transcript",
    "subtitle_srt": "transcript",
    "local_csv": "table",
    "local_tsv": "table",
    "local_xlsx": "table",
    "local_spreadsheet_unstructured": "table",
}


def finalize_source_meta(result: "DoclingParseResult", filename: str | None) -> None:
    """M2 + routing_trace finalizer — called once at the worker boundary.

    Fills fallbacks (title ← cleaned filename stem), a deterministic
    format-family source_type (semantic refinement is Ghost A's job later),
    and the per-document routing_trace (every cascade decision). Idempotent.

    T-HOOK-3: also runs the date de-conflation resolver over the parser's
    typed ``date_candidates`` — ``document_date``/``source_published_at`` are
    set ONLY from publication-grade candidates (file-creation/revision times
    never leak in; unknown stays null with a reason code) — and stamps the
    full bibliographic block onto ``routing_trace["bibliographic"]``, which
    ``mongo_writer.upsert_document`` promotes to top-level document fields.
    """
    fname = filename or result.filename or ""
    stem = Path(fname).stem
    stem_title = re.sub(r"[_\-]+", " ", stem).strip()[:300] or None
    # stable across repeat calls: a title equal to the stem fallback counts as
    # filename-derived even if set by a prior finalize pass (idempotency)
    title_from_metadata = bool(result.title) and result.title != stem_title
    if not result.title:
        result.title = stem_title
    if not result.source_type:
        fmt = result.source_format or ""
        if fmt.startswith("code_"):
            result.source_type = "code"
        else:
            result.source_type = _SOURCE_TYPE_BY_FORMAT.get(fmt, "document")
    tier = getattr(result.source_tier, "value", str(result.source_tier))
    parent_strategy = {
        "ocr_ast": "pdf_page_grouped",
        "tier_code": "ast_bound_code",
        "tier_c": "semantic_parents_or_token_window",
    }.get(tier, "heading_bound")
    # T-HOOK-3 — resolve typed date candidates; publication vs file-time
    # de-conflation lives in services.ingestion.bibliographic.
    resolution = resolve_document_dates(getattr(result, "date_candidates", None) or [])
    result.document_date = resolution["document_date"]
    prior_block = (
        (getattr(result, "routing_trace", None) or {}).get("bibliographic") or {}
    )
    prior_provenance = prior_block.get("bibliographic_provenance") or {}
    result.routing_trace = {
        "parser": result.source_format or "docling_sidecar",
        "tier": tier,
        "has_structure": bool(result.has_structure),
        "augmented_headers": bool(result.augmented_with_synthetic_headers),
        "language": result.language,
        "num_pages": result.num_pages,
        "parent_strategy": parent_strategy,
        "child_strategy": "semantic_split+routers",
        "title_source": "metadata" if title_from_metadata else "filename",
        # promoted to top-level document fields at the Mongo writer boundary
        "bibliographic": {
            "title": result.title,
            "author": result.author,
            "language": getattr(result, "language_meta", None),
            "document_date": resolution["document_date"],
            "source_published_at": resolution["source_published_at"],
            "date_confidence": resolution["date_confidence"],
            "bibliographic_provenance": build_provenance(
                method=resolution["method"],
                source=resolution["source"],
                precision=resolution["precision"],
                reason=resolution["reason"],
                origin="ingest",
                captured_at=prior_provenance.get("captured_at"),
            ),
        },
    }


def _strip_yaml_frontmatter(text: str) -> str:
    if not text or not text.lstrip("﻿").startswith("---"):
        return text
    stripped = _FRONTMATTER_RE.sub("", text, count=1)
    if stripped is not text:
        logger.info("local_markdown: stripped YAML frontmatter (%d chars)",
                    len(text) - len(stripped))
    return stripped


# Bold-key metadata line: `**Source:** https://…`, `**Extracted:** 2026-03-24`.
# Scraper/export tooling writes a run of these right under the document title —
# a second metadata layer alongside YAML frontmatter. Like frontmatter, it is
# provenance, not content: it embeds as a junk chunk and its dates/URLs sit one
# entity-mention away from becoming junk facts.
_BOLD_KEY_LINE = re.compile(r"^\*\*[^*\n]{1,32}:\*\*\s")
_MD_TITLE = re.compile(r"^#{1,6}\s")


def _strip_leading_metadata_block(text: str) -> str:
    """Drop a doc-START run of 2+ bold-key metadata lines (blank lines allowed
    between them), preserving an optional leading title heading. A single bold
    note is left alone — only a BLOCK signals export metadata."""
    if not text:
        return text
    lines = text.split("\n")
    i = 0
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines) and _MD_TITLE.match(lines[i].strip()):
        i += 1
    j, n_meta = i, 0
    while j < len(lines):
        s = lines[j].strip()
        if not s:
            j += 1
            continue
        if _BOLD_KEY_LINE.match(s):
            n_meta += 1
            j += 1
            continue
        break
    if n_meta < 2:
        return text
    logger.info("local_markdown: stripped %d leading metadata lines", n_meta)
    return "\n".join(lines[:i] + [""] + lines[j:])


def _looks_like_html(filename: str, mime: str) -> bool:
    return (mime or "").lower() in _HTML_MIMES or _extension(filename) in _HTML_EXTS


def _looks_like_epub(filename: str, mime: str) -> bool:
    return (mime or "").lower() in _EPUB_MIMES or _extension(filename) in _EPUB_EXTS


def _looks_like_docx(filename: str, mime: str) -> bool:
    return (mime or "").lower() in _DOCX_MIMES or _extension(filename) == ".docx"


def _looks_like_csv(filename: str, mime: str) -> bool:
    return (mime or "").lower() in _CSV_MIMES or _extension(filename) in _CSV_EXTS


def _looks_like_spreadsheet(filename: str, mime: str) -> bool:
    return (mime or "").lower() in _SPREADSHEET_MIMES or _extension(filename) in _SPREADSHEET_EXTS


def _looks_like_code(filename: str) -> bool:
    basename = Path(filename or "").name.lower()
    return basename in _CODE_FILENAME_TO_LANGUAGE or _extension(filename) in _CODE_EXT_TO_LANGUAGE


def docling_sidecar_needed(filename: str, mime: str) -> bool:
    """True only for formats that cannot use the local parser path."""
    if _looks_like_code(filename):
        return False
    if _looks_like_pdf(filename, mime):
        return False
    if _looks_like_markdown(filename, mime):
        return False
    if _looks_like_html(filename, mime):
        return False
    if _looks_like_epub(filename, mime):
        return False
    if _looks_like_docx(filename, mime):
        return False
    if _looks_like_csv(filename, mime) or _looks_like_spreadsheet(filename, mime):
        return False
    if _looks_like_plain_text(filename, mime):
        return False
    return True


def parser_strategy(filename: str, mime: str) -> str:
    """Human-readable parse lane for diagnostics and tests."""
    if _looks_like_code(filename):
        return "local_code"
    if _looks_like_pdf(filename, mime):
        return "local_pdf_fast_text"
    if _looks_like_markdown(filename, mime):
        return "local_markdown"
    if _looks_like_html(filename, mime):
        return "local_html"
    if _looks_like_epub(filename, mime):
        return "local_epub"
    if _looks_like_docx(filename, mime):
        return "local_docx"
    if _looks_like_csv(filename, mime):
        return "local_csv"
    if _looks_like_spreadsheet(filename, mime):
        return "local_spreadsheet"
    if _looks_like_plain_text(filename, mime):
        return "local_text"
    return "docling_sidecar"


@dataclass
class Section:
    heading_path: list[str]
    text: str
    element_type: str  # "section_heading" | "paragraph" | "code_block" | ...
    level: int | None = None
    # Code lane: language tag (e.g. "python", "luau"); None for prose sections.
    language: str | None = None
    # Structured element metadata. Used by local markdown/table parsing while
    # keeping sidecar-produced prose/code sections backward-compatible.
    metadata: dict = field(default_factory=dict)


@dataclass
class DoclingParseResult:
    text: str
    markdown: str
    sections: list[Section]
    pages: list[str] | None
    has_structure: bool
    source_tier: SourceTier
    h1_count: int = 0
    h2_count: int = 0
    num_pages: int = 1
    source_format: str = ""
    augmented_with_synthetic_headers: bool = False
    injected_headers_audit: list[dict] = field(default_factory=list)
    # Code lane: filled by the early-intercept gate for code files (.py/.rs/etc).
    # None for prose / markdown / PDF / EPUB ingest.
    language: str | None = None
    # Original upload filename — needed by tier_chunker to stamp file_path
    # into per-chunk metadata for the code lane.
    filename: str | None = None
    # M2 — parse-time source metadata (POLYMATH_ARCHITECTURE §2.2 / §3.S1).
    # Filled by each parser from its format (PDF info, DOCX core-props, MD
    # frontmatter, HTML meta); filename-stem fallback applied downstream.
    # Prerequisite for two-lane anchoring + the summary-tree compact schema.
    title: str | None = None
    author: str | None = None
    document_date: str | None = None  # PUBLICATION date only (resolved in finalize)
    source_type: str | None = None    # book|paper|standard|manual|blog|... (heuristic)
    # T-HOOK-3 — bibliographic/date identity. `language_meta` is the natural
    # language from EXPLICIT format metadata (frontmatter lang, <html lang>,
    # EPUB dc:language, DOCX core-props) — distinct from the code-lane
    # `language` (programming language). `date_candidates` collects every raw
    # dated observation (typed publication vs file-creation vs revision);
    # `finalize_source_meta` runs the de-conflation resolver over them.
    language_meta: str | None = None
    date_candidates: list = field(default_factory=list)
    # Per-document routing report — cascade decisions (intercept → sniff →
    # tier → parent strategy) recorded for /documents/{id} visibility.
    routing_trace: dict = field(default_factory=dict)


def _looks_like_plain_text(filename: str, mime: str) -> bool:
    ext = _extension(filename)
    mime_l = (mime or "").lower()
    if mime_l in _PLAIN_TEXT_MIMES:
        return True
    if ext in _PLAIN_TEXT_EXTS:
        return True
    # Browsers sometimes upload markdown/text as octet-stream. Accept only
    # known text-like extensions here so office files do not get decoded as
    # mojibake just because the client supplied a vague MIME.
    return mime_l == "application/octet-stream" and ext in (_PLAIN_TEXT_EXTS | _MARKDOWN_EXTS)


_FENCE_OPEN_RE = re.compile(r"^```([a-zA-Z0-9_+\-]*)\s*$")
_FENCE_CLOSE_RE = re.compile(r"^```\s*$")
_HEADING_ANCHOR_RE = re.compile(r"\s*\{#[^\n}]*\}\s*$")
_TABLE_SEPARATOR_CELL_RE = re.compile(r"^:?-{3,}:?$")
_TABLE_CAPTION_RE = re.compile(
    r"^(?:table|tbl\.?)\s*[\w.\-]+(?:\s*[:.\-]\s+|\s+).+",
    re.IGNORECASE,
)


def _split_markdown_table_row(line: str) -> list[str]:
    """Split a GitHub-style pipe table row into cleaned cells."""
    row = (line or "").strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [
        cell.replace(r"\|", "|").strip()
        for cell in re.split(r"(?<!\\)\|", row)
    ]


# Inline markdown link `[text](url)` / `[text](url "title")` — keep the
# visible text, drop the target. Scraped docs carry anchor links on every
# heading (`Getting Started[¶](https://…#x "Link to this heading")`) and the
# URLs/pilcrows were the last embedding-noise class left after the
# frontmatter/metadata strips. Applied to prose paragraphs, heading titles,
# and table cells — NEVER to code blocks (they're sectioned separately and
# must stay verbatim).
_IMG_MD_RE = re.compile(r"!\[[^\]\n]*\]\([^)\n]*\)")  # drop images entirely (alt incl.)
_INLINE_MD_LINK_RE = re.compile(r"\[([^\]\n]*)\]\((?:[^)\s]+(?:\s+\"[^\"]*\")?)\)")
_EMPTY_MD_ANCHOR_RE = re.compile(r"\[\]\{#[^}\n]+\}")
_BARE_URL_RE = re.compile(r"https?://\S+")
_TRANSCRIPT_HEADER_RE = re.compile(
    r"^\s*(Video|URL|Duration|Segments|Source|Date)\s*:\s*(.*?)\s*$",
    re.IGNORECASE,
)
_TRANSCRIPT_SEGMENT_RE = re.compile(
    r"^\s*(?:\[(?P<bracket_time>(?:\d+:)?\d{1,2}:\d{2}(?:[.,]\d{1,3})?)\]"
    r"|(?P<plain_time>(?:\d+:)?\d{1,2}:\d{2}(?:[.,]\d{1,3})?))"
    r"\s*(?:[-–—]\s*)?(?P<text>.+?)\s*$"
)


def _scrub_inline_links(text: str) -> str:
    if not text or ("[" not in text and "http" not in text and "¶" not in text):
        return text
    t = _IMG_MD_RE.sub("", text)  # before link rewrite, or `![x](u)` leaves a stray `!`
    t = _EMPTY_MD_ANCHOR_RE.sub("", t)
    t = _INLINE_MD_LINK_RE.sub(r"\1", t)
    t = _BARE_URL_RE.sub("", t)
    t = t.replace("¶", "")
    return re.sub(r"[ \t]{2,}", " ", t)


def retrievable_content_text(result: "DoclingParseResult") -> str:
    """Return visible source content after deterministic markup removal.

    This is a content-presence gate, not a relevance heuristic: any visible
    letter or number keeps the source ingestible. It only rejects artifacts
    made entirely of images, empty anchors, URLs, and markup.
    """

    section_text = "\n".join(
        str(section.text or "")
        for section in (result.sections or [])
        if str(section.text or "").strip()
    )
    candidate = section_text or str(result.text or result.markdown or "")
    if str(result.source_format or "") == "local_markdown":
        candidate = _scrub_inline_links(candidate)
        candidate = re.sub(r"<[^>]+>", " ", candidate)
        candidate = re.sub(r"(?m)^\s*[#>*_`~-]+\s*$", " ", candidate)
    return " ".join(candidate.split()).strip()


def has_retrievable_content(result: "DoclingParseResult") -> bool:
    """Whether a parsed source contains any visible semantic token."""

    return bool(re.search(r"[^\W_]{2,}", retrievable_content_text(result), re.UNICODE))


_SUB_TIME_RE = re.compile(
    r"(?:(\d{1,2}):)?(\d{1,2}):(\d{2})[.,](\d{3})\s*-->\s*(?:(\d{1,2}):)?(\d{1,2}):(\d{2})[.,](\d{3})"
)
_SUB_SPEAKER_VTT_RE = re.compile(r"<v\s+([^>]+)>")
_SUB_SPEAKER_LINE_RE = re.compile(r"^([A-Z][\w .'-]{0,24}):\s+")
_SUB_TAG_RE = re.compile(r"<[^>]+>")


def _parse_subtitle_file(raw_bytes: bytes, filename: str, mime: str):
    """Router 4 (POLYMATH_ARCHITECTURE §3.S2): VTT/SRT subtitle transcripts.

    Stdlib cue parsing → the SAME transcript_block sections the YouTube path
    emits (time ranges + speakers in metadata, semantic_split guarded off via
    source_format), so retrieval/hydration treat all transcripts identically.
    Returns None when the file is not a subtitle (falls through to other lanes).
    """
    ext = _extension(filename)
    if ext not in _SUBTITLE_EXTS and (mime or "").lower() not in ("text/vtt",):
        return None
    try:
        text = raw_bytes.decode("utf-8", errors="replace")
    except Exception:
        return None
    is_vtt = ext == ".vtt" or text.lstrip()[:6].upper().startswith("WEBVTT")

    cues: list[tuple[str, str, str, str]] = []  # (start, end, speaker, text)
    cur_time: tuple[str, str] | None = None
    cur_lines: list[str] = []

    def _flush_cue():
        nonlocal cur_time, cur_lines
        if cur_time and cur_lines:
            body = " ".join(ln.strip() for ln in cur_lines if ln.strip())
            speaker = ""
            m = _SUB_SPEAKER_VTT_RE.search(body)
            if m:
                speaker = m.group(1).strip()
            body = _SUB_TAG_RE.sub("", body).strip()
            if not speaker:
                m2 = _SUB_SPEAKER_LINE_RE.match(body)
                if m2:
                    speaker = m2.group(1).strip()
                    body = body[m2.end():].strip()
            if body:
                cues.append((cur_time[0], cur_time[1], speaker, body))
        cur_time, cur_lines = None, []

    for line in text.splitlines():
        s = line.strip()
        if not s:
            _flush_cue()
            continue
        tm = _SUB_TIME_RE.search(s)
        if tm:
            _flush_cue()

            def _fmt(h, m, sec):
                return f"{h}:{m}:{sec}" if h else f"{m}:{sec}"

            cur_time = (
                _fmt(tm.group(1), tm.group(2), tm.group(3)),
                _fmt(tm.group(5), tm.group(6), tm.group(7)),
            )
            continue
        if cur_time is None:
            continue  # WEBVTT header, cue ids, NOTE/STYLE blocks
        cur_lines.append(s)
    _flush_cue()

    if len(cues) < 3:
        return None  # not a real subtitle file — let other lanes try

    title = Path(filename or "subtitle").stem or "subtitle"
    fmt = "subtitle_vtt" if is_vtt else "subtitle_srt"
    sections: list[Section] = []
    group: list[tuple[str, str, str, str]] = []
    words = 0
    start_idx = 0

    def _emit_group(first_idx: int):
        nonlocal group, words
        if not group:
            return
        speakers = sorted({c[2] for c in group if c[2]})
        lines = [(f"{c[2]}: {c[3]}" if c[2] else c[3]) for c in group]
        sections.append(
            Section(
                heading_path=[title],
                text=f"[{group[0][0]} - {group[-1][1]}]\n" + "\n".join(lines),
                element_type="transcript_block",
                metadata={
                    "source_format": fmt,
                    "video_title": title,
                    "time_start": group[0][0],
                    "time_end": group[-1][1],
                    "segment_start": first_idx,
                    "segment_end": first_idx + len(group) - 1,
                    "speakers": speakers,
                },
            )
        )
        group, words = [], 0

    for i, cue in enumerate(cues):
        w = len(cue[3].split())
        if group and words + w > 120:
            _emit_group(start_idx)
            start_idx = i
        group.append(cue)
        words += w
    _emit_group(start_idx)

    return DoclingParseResult(
        text="\n\n".join(s.text for s in sections),
        markdown="\n\n".join(s.text for s in sections),
        sections=sections,
        pages=None,
        has_structure=True,
        source_tier=SourceTier.tier_b,
        h1_count=1,
        h2_count=0,
        source_format=fmt,
        filename=filename,
        title=title,
    )


def _parse_transcript_text_document(text: str, filename: str) -> DoclingParseResult | None:
    """Parse timestamped transcript exports into bounded semantic blocks.

    Expected shape is deterministic YouTube/transcript text:

        Video: ...
        URL: ...
        Duration: ...
        Segments: ...
        Source: YouTube Transcript API
        Date: ...

        [0:00] speech text
        [0:03] more speech text

    Header fields become chunk metadata; timestamped speech becomes the only
    substantive retrieval text. This prevents the first child chunk from being
    mostly `URL`/`Duration`/`Segments` metadata while preserving time ranges
    for source display and later citation jumps.
    """
    if not text or not re.search(r"(?:^|\n)\s*(?:\[?(?:\d+:)?\d{1,2}:\d{2})", text):
        return None

    headers: dict[str, str] = {}
    segments: list[dict[str, str]] = []
    seen_segment = False
    content_lines = 0

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        content_lines += 1

        match = _TRANSCRIPT_SEGMENT_RE.match(line)
        if match:
            seen_segment = True
            spoken = re.sub(r"\s+", " ", match.group("text")).strip()
            if spoken:
                segments.append(
                    {
                        "time": (
                            match.group("bracket_time") or match.group("plain_time")
                        ).replace(",", "."),
                        "text": spoken,
                    }
                )
            continue

        header = _TRANSCRIPT_HEADER_RE.match(line)
        if header and not seen_segment:
            key = header.group(1).lower()
            value = re.sub(r"\s+", " ", header.group(2)).strip()
            if value:
                headers[key] = value
            continue

        # Transcript continuation line. Some transcript tools wrap long
        # captions without repeating a timestamp. Attach those words to the
        # previous timed segment instead of creating an untimed junk paragraph.
        if seen_segment and segments:
            continuation = re.sub(r"\s+", " ", line).strip()
            segments[-1]["text"] = f"{segments[-1]['text']} {continuation}".strip()

    source_hint = f" {headers.get('source', '')} {headers.get('url', '')} {filename} ".lower()
    timestamp_density = len(segments) / max(content_lines, 1)
    looks_like_transcript = (
        len(segments) >= 3
        and (
            "youtube transcript" in source_hint
            or "youtu.be" in source_hint
            or "youtube.com" in source_hint
            or "segments" in headers
            or "duration" in headers
            or (len(segments) >= 5 and timestamp_density >= 0.60)
        )
    )
    if not looks_like_transcript:
        return None

    title = headers.get("video") or filename or "Transcript"
    max_words_per_block = 120
    grouped: list[tuple[int, int, list[dict[str, str]]]] = []
    start = 0
    buf: list[dict[str, str]] = []
    word_count = 0
    for idx, segment in enumerate(segments):
        words = len(segment["text"].split())
        if buf and word_count + words > max_words_per_block:
            grouped.append((start, idx - 1, buf))
            start = idx
            buf = []
            word_count = 0
        buf.append(segment)
        word_count += words
    if buf:
        grouped.append((start, start + len(buf) - 1, buf))

    sections: list[Section] = [
        Section(
            heading_path=[title],
            text=title,
            element_type="section_heading",
            level=1,
        )
    ]

    for start_idx, end_idx, group in grouped:
        time_start = group[0]["time"]
        time_end = group[-1]["time"]
        speech = " ".join(item["text"] for item in group)
        block_text = (
            f"Video: {title}\n"
            f"Transcript range: {time_start}-{time_end}\n\n"
            f"{speech}"
        )
        sections.append(
            Section(
                heading_path=[title],
                text=block_text,
                element_type="transcript_block",
                metadata={
                    "source_format": "youtube_transcript",
                    "video_title": title,
                    "url": headers.get("url", ""),
                    "duration": headers.get("duration", ""),
                    "segments_declared": headers.get("segments", ""),
                    "source": headers.get("source", ""),
                    "date": headers.get("date", ""),
                    "time_start": time_start,
                    "time_end": time_end,
                    "segment_start": start_idx,
                    "segment_end": end_idx,
                },
            )
        )

    return DoclingParseResult(
        text="\n\n".join(section.text for section in sections),
        markdown="\n\n".join(section.text for section in sections),
        sections=sections,
        pages=None,
        has_structure=True,
        source_tier=SourceTier.tier_b,
        h1_count=1,
        h2_count=0,
        source_format="youtube_transcript",
        filename=filename,
        title=title,
    )


def _clean_heading_title(title: str) -> str:
    """Strip markdown anchor IDs + inline links from visible heading metadata."""
    return _scrub_inline_links(_HEADING_ANCHOR_RE.sub("", title or "")).strip()


def _is_markdown_table_separator(line: str) -> bool:
    cells = _split_markdown_table_row(line)
    if len(cells) < 2:
        return False
    return all(_TABLE_SEPARATOR_CELL_RE.match(cell.strip()) for cell in cells)


def _is_markdown_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    header = lines[index]
    separator = lines[index + 1]
    if "|" not in header or "|" not in separator:
        return False
    header_cells = _split_markdown_table_row(header)
    if len([cell for cell in header_cells if cell.strip()]) < 2:
        return False
    if not _is_markdown_table_separator(separator):
        return False
    separator_cells = _split_markdown_table_row(separator)
    return len(separator_cells) == len(header_cells)


def _consume_markdown_table(
    lines: list[str],
    index: int,
) -> tuple[list[str], list[str], list[list[str]], int]:
    """Return raw table lines, columns, rows, and next index."""
    raw_lines = [lines[index], lines[index + 1]]
    columns = _split_markdown_table_row(lines[index])
    rows: list[list[str]] = []
    j = index + 2
    while j < len(lines):
        row_line = lines[j]
        if not row_line.strip() or "|" not in row_line:
            break
        cells = _split_markdown_table_row(row_line)
        if len(cells) < 2:
            break
        raw_lines.append(row_line)
        rows.append(cells)
        j += 1
    return raw_lines, columns, rows, j


def _pop_table_caption(paragraph: list[str]) -> str:
    """Detach an immediate `Table N. ...` caption from the prose buffer."""
    if not paragraph:
        return ""
    i = len(paragraph) - 1
    while i >= 0 and not paragraph[i].strip():
        i -= 1
    if i < 0:
        return ""
    candidate = paragraph[i].strip()
    if len(candidate) > 240 or not _TABLE_CAPTION_RE.match(candidate):
        return ""
    del paragraph[i:]
    while paragraph and not paragraph[-1].strip():
        paragraph.pop()
    return candidate


def _table_label(caption: str, table_index: int) -> str:
    if caption:
        match = re.match(r"^((?:table|tbl\.?)\s*[\w.\-]+)", caption, re.IGNORECASE)
        if match:
            return match.group(1).replace("tbl.", "Table").strip()
    return f"Table {table_index}"


def _linearize_markdown_table(
    *,
    heading_path: list[str],
    caption: str,
    table_index: int,
    columns: list[str],
    rows: list[list[str]],
) -> str:
    """Render a markdown table as row-wise text for embeddings and extraction."""
    cleaned_columns = [
        re.sub(r"\s+", " ", col).strip() or f"column_{i + 1}"
        for i, col in enumerate(columns)
    ]
    lines: list[str] = [f"Table: {_table_label(caption, table_index)}"]
    if heading_path:
        lines.append(f"Section: {' > '.join(heading_path)}")
    if caption:
        lines.append(f"Caption: {caption}")
    lines.append(f"Columns: {' | '.join(cleaned_columns)}")
    lines.append("")

    for row_idx, row in enumerate(rows, start=1):
        padded = list(row[: len(cleaned_columns)])
        if len(padded) < len(cleaned_columns):
            padded.extend([""] * (len(cleaned_columns) - len(padded)))
        pairs = []
        for col, cell in zip(cleaned_columns, padded):
            value = _scrub_inline_links(re.sub(r"\s+", " ", cell)).strip()
            if value:
                pairs.append(f"{col}={value}")
        if pairs:
            lines.append(f"Row {row_idx}: " + "; ".join(pairs))

    return "\n".join(lines).strip()


def _clean_table_cell(value: object) -> str:
    if value is None:
        return ""
    text = str(value)
    return _scrub_inline_links(re.sub(r"\s+", " ", text)).strip()


def _table_columns_and_rows(
    rows: list[list[str]],
    *,
    has_header: bool = True,
) -> tuple[list[str], list[list[str]]]:
    cleaned_rows = [
        [_clean_table_cell(cell) for cell in row]
        for row in rows
        if any(_clean_table_cell(cell) for cell in row)
    ]
    if not cleaned_rows:
        return [], []

    width = max(len(row) for row in cleaned_rows)
    padded = [row + [""] * (width - len(row)) for row in cleaned_rows]

    if has_header and padded:
        raw_columns = padded[0]
        data_rows = padded[1:]
    else:
        raw_columns = []
        data_rows = padded

    seen: dict[str, int] = {}
    columns: list[str] = []
    for idx in range(width):
        name = _clean_table_cell(raw_columns[idx] if idx < len(raw_columns) else "")
        if not name:
            name = f"column_{idx + 1}"
        count = seen.get(name.lower(), 0) + 1
        seen[name.lower()] = count
        columns.append(name if count == 1 else f"{name}_{count}")
    return columns, data_rows


def _tabular_result(
    *,
    filename: str,
    source_format: str,
    table_sections: list[Section],
    h1_count: int = 1,
) -> DoclingParseResult:
    title = filename or "Table Upload"
    sections: list[Section] = [
        Section(
            heading_path=[title],
            text=title,
            element_type="section_heading",
            level=1,
        )
    ]
    sections.extend(table_sections)
    text = "\n\n".join(section.text for section in sections if section.text)
    return DoclingParseResult(
        text=text,
        markdown=text,
        sections=sections,
        pages=None,
        has_structure=bool(table_sections),
        source_tier=SourceTier.tier_b if table_sections else SourceTier.tier_c,
        h1_count=h1_count if table_sections else 0,
        h2_count=0,
        source_format=source_format,
        filename=filename,
    )


def _parse_delimited_table_document(
    raw_bytes: bytes,
    filename: str,
    mime: str,
) -> DoclingParseResult | None:
    if not _looks_like_csv(filename, mime):
        return None

    text = raw_bytes.decode("utf-8-sig", errors="replace")
    ext = _extension(filename)
    delimiter = "\t" if ext == ".tsv" or (mime or "").lower() == "text/tab-separated-values" else ","
    try:
        sniffed = csv.Sniffer().sniff(text[:8192], delimiters=",\t;|")
        delimiter = sniffed.delimiter
    except Exception:
        pass
    try:
        has_header = csv.Sniffer().has_header(text[:8192])
    except Exception:
        has_header = True

    rows = list(csv.reader(StringIO(text), delimiter=delimiter))
    columns, data_rows = _table_columns_and_rows(rows, has_header=has_header)
    if not columns and not data_rows:
        return None

    source_format = "local_tsv" if delimiter == "\t" else "local_csv"
    table_text = _linearize_markdown_table(
        heading_path=[filename] if filename else [],
        caption=f"{source_format.replace('_', ' ').title()} {filename}".strip(),
        table_index=1,
        columns=columns,
        rows=data_rows,
    )
    table = Section(
        heading_path=[filename] if filename else [],
        text=table_text,
        element_type="table",
        metadata={
            "table_index": 1,
            "caption": filename,
            "columns": columns,
            "row_count": len(data_rows),
            "delimiter": delimiter,
            "source_format": source_format,
        },
    )
    return _tabular_result(
        filename=filename,
        source_format=source_format,
        table_sections=[table],
    )


def _parse_xlsx_table_document(
    raw_bytes: bytes,
    filename: str,
    mime: str,
) -> DoclingParseResult | None:
    ext = _extension(filename)
    if not _looks_like_spreadsheet(filename, mime) or ext == ".xls":
        return None
    try:
        from openpyxl import load_workbook
    except Exception:
        return None

    try:
        workbook = load_workbook(BytesIO(raw_bytes), read_only=True, data_only=True)
    except Exception as exc:
        logger.warning("openpyxl failed for %s: %s", filename, exc)
        return None

    table_sections: list[Section] = []
    for table_index, sheet in enumerate(workbook.worksheets[:_TABLE_PARSE_MAX_SHEETS], start=1):
        rows: list[list[str]] = []
        truncated = False
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if row_idx > _TABLE_PARSE_MAX_ROWS_PER_SHEET:
                truncated = True
                break
            cleaned = [_clean_table_cell(cell) for cell in row]
            if any(cleaned):
                rows.append(cleaned)
        columns, data_rows = _table_columns_and_rows(rows, has_header=True)
        if not columns and not data_rows:
            continue
        heading_path = [filename, sheet.title] if filename else [sheet.title]
        table_text = _linearize_markdown_table(
            heading_path=heading_path,
            caption=f"Sheet {sheet.title}",
            table_index=table_index,
            columns=columns,
            rows=data_rows,
        )
        table_sections.append(
            Section(
                heading_path=heading_path,
                text=table_text,
                element_type="table",
                metadata={
                    "table_index": table_index,
                    "caption": f"Sheet {sheet.title}",
                    "sheet_name": sheet.title,
                    "columns": columns,
                    "row_count": len(data_rows),
                    "source_format": "local_xlsx",
                    "truncated": truncated,
                    "max_rows_per_sheet": _TABLE_PARSE_MAX_ROWS_PER_SHEET,
                },
            )
        )

    try:
        workbook.close()
    except Exception:
        pass

    if not table_sections:
        return None
    return _tabular_result(
        filename=filename,
        source_format="local_xlsx",
        table_sections=table_sections,
    )


def _parse_legacy_spreadsheet_document(
    raw_bytes: bytes,
    filename: str,
    mime: str,
) -> DoclingParseResult | None:
    if not _looks_like_spreadsheet(filename, mime):
        return None
    from services.ingestion.format_router import route

    decoded = route(raw_bytes, filename=filename, mime_hint=mime)
    text = (decoded.text or "").strip()
    if not text:
        return None
    section = Section(
        heading_path=[filename] if filename else [],
        text=text,
        element_type="table",
        metadata={
            "table_index": 1,
            "caption": filename,
            "source_format": "local_spreadsheet_unstructured",
        },
    )
    return _tabular_result(
        filename=filename,
        source_format="local_spreadsheet_unstructured",
        table_sections=[section],
    )


def _markdown_sections(markdown: str) -> tuple[list[Section], int, int]:
    """Local Markdown section walker, code-fence-aware.

    Headings define the hierarchy; paragraph text inherits the active heading
    path. Triple-backtick fenced code blocks are emitted as separate
    Section(element_type="code_block", language=<tag>) entries so the code lane
    in tier_chunker can route them through code_splitter.pack() instead of the
    prose sentence/token splitters. Tilde fences (~~~) and indented fences fall
    through as prose — Phase 1 limitation, documented for later expansion.
    """
    sections: list[Section] = []
    heading_stack: list[str] = []
    current_path: list[str] = []
    paragraph: list[str] = []
    h1_count = 0
    h2_count = 0
    table_count = 0

    def flush_paragraph() -> None:
        nonlocal paragraph
        text = _scrub_inline_links("\n".join(paragraph)).strip()
        if text:
            sections.append(
                Section(
                    heading_path=list(current_path),
                    text=text,
                    element_type="paragraph",
                )
            )
        paragraph = []

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fence open? Consume until matching close (or EOF).
        fence_open = _FENCE_OPEN_RE.match(line.strip())
        if fence_open is not None:
            flush_paragraph()
            language = (fence_open.group(1) or "").lower()
            fence_lines: list[str] = [line]
            j = i + 1
            while j < len(lines):
                fence_lines.append(lines[j])
                if _FENCE_CLOSE_RE.match(lines[j].strip()):
                    j += 1
                    break
                j += 1
            sections.append(
                Section(
                    heading_path=list(current_path),
                    text="\n".join(fence_lines),
                    element_type="code_block",
                    language=language or None,
                )
            )
            i = j
            continue

        if _is_markdown_table_start(lines, i):
            caption = _pop_table_caption(paragraph)
            flush_paragraph()
            table_count += 1
            _raw_lines, columns, rows, j = _consume_markdown_table(lines, i)
            metadata = {
                "table_index": table_count,
                "caption": caption,
                "columns": [
                    re.sub(r"\s+", " ", col).strip() or f"column_{idx + 1}"
                    for idx, col in enumerate(columns)
                ],
                "row_count": len(rows),
                "source_format": "markdown_pipe_table",
            }
            sections.append(
                Section(
                    heading_path=list(current_path),
                    text=_linearize_markdown_table(
                        heading_path=list(current_path),
                        caption=caption,
                        table_index=table_count,
                        columns=columns,
                        rows=rows,
                    ),
                    element_type="table",
                    metadata=metadata,
                )
            )
            i = j
            continue

        match = re.match(r"^(#{1,6})\s+(.+?)\s*#*\s*$", line.strip())
        if not match:
            paragraph.append(line)
            i += 1
            continue

        flush_paragraph()
        level = len(match.group(1))
        title = _clean_heading_title(match.group(2).strip())
        if not title:
            i += 1
            continue
        if level == 1:
            h1_count += 1
        elif level == 2:
            h2_count += 1
        heading_stack = heading_stack[: level - 1]
        while len(heading_stack) < level - 1:
            heading_stack.append("")
        heading_stack.append(title)
        current_path = [part for part in heading_stack if part]
        sections.append(
            Section(
                heading_path=list(current_path),
                text=title,
                element_type="section_heading",
                level=level,
            )
        )
        i += 1

    flush_paragraph()
    return sections, h1_count, h2_count


def _parse_local_text_document(raw_bytes: bytes, filename: str, mime: str) -> DoclingParseResult | None:
    """Parse text/markdown/html without the Docling sidecar.

    Local parsing keeps default Docker startup API-first. Docling remains an
    explicit profile for formats that truly need layout-aware conversion.
    """
    csv_result = _parse_delimited_table_document(raw_bytes, filename, mime)
    if csv_result is not None:
        return csv_result

    xlsx_result = _parse_xlsx_table_document(raw_bytes, filename, mime)
    if xlsx_result is not None:
        return xlsx_result

    legacy_sheet_result = _parse_legacy_spreadsheet_document(raw_bytes, filename, mime)
    if legacy_sheet_result is not None:
        return legacy_sheet_result

    if _looks_like_epub(filename, mime):
        epub_result = _parse_local_epub_document(raw_bytes, filename)
        if epub_result is not None:
            return epub_result

    if _looks_like_html(filename, mime):
        from services.ingestion.format_router import route

        decoded = route(raw_bytes, filename=filename, mime_hint=mime)
        text = decoded.text or ""
        html_result = DoclingParseResult(
            text=text,
            markdown=text,
            sections=[],
            pages=None,
            has_structure=False,
            source_tier=SourceTier.tier_b,
            source_format="local_html",
            filename=filename,
        )
        _apply_meta(html_result, _meta_from_html(raw_bytes))  # M2 + T-HOOK-3
        return html_result

    if _looks_like_markdown(filename, mime):
        raw_text = raw_bytes.decode("utf-8", errors="replace")
        fm_meta = _meta_from_frontmatter(raw_text)  # M2: read before stripping
        fm_candidates = list(fm_meta.get("date_candidates") or [])
        markdown = _strip_leading_metadata_block(_strip_yaml_frontmatter(raw_text))
        sections, h1, h2 = _markdown_sections(markdown)
        has_tables = any(s.element_type == "table" for s in sections)
        has_structure = (h1 + h2) > 0 or has_tables
        source_tier = (
            SourceTier.tier_a
            if (h1 + h2) > 0
            else SourceTier.tier_b if has_tables else SourceTier.tier_c
        )
        return DoclingParseResult(
            text=markdown,
            markdown=markdown,
            sections=sections,
            pages=None,
            has_structure=has_structure,
            source_tier=source_tier,
            h1_count=h1,
            h2_count=h2,
            source_format="local_markdown",
            filename=filename,
            title=fm_meta.get("title"),
            author=fm_meta.get("author"),
            document_date=resolve_document_dates(fm_candidates)["document_date"],
            language_meta=fm_meta.get("language_meta"),
            date_candidates=fm_candidates,
        )

    if _looks_like_docx(filename, mime):
        docx_result = _parse_local_docx_document(raw_bytes, filename)
        if docx_result is not None:
            _apply_meta(docx_result, _meta_from_docx(raw_bytes))  # M2 core-props
            return docx_result

    if not _looks_like_plain_text(filename, mime):
        return None

    raw_text = raw_bytes.decode("utf-8", errors="replace")
    transcript_result = _parse_transcript_text_document(raw_text, filename)
    if transcript_result is not None:
        return transcript_result

    aug_bytes, aug_filename, aug_mime, augmented, audit = _maybe_augment_plaintext(
        raw_bytes, filename, mime
    )
    text = aug_bytes.decode("utf-8", errors="replace")
    sections, h1, h2 = _markdown_sections(text)
    has_tables = any(s.element_type == "table" for s in sections)
    has_structure = (augmented and (h1 + h2) > 0) or has_tables
    source_tier = (
        SourceTier.tier_b_plus
        if augmented and (h1 + h2) > 0
        else SourceTier.tier_b if has_tables else SourceTier.tier_c
    )
    return DoclingParseResult(
        text=text,
        markdown=text,
        sections=sections,
        pages=None,
        has_structure=has_structure,
        source_tier=source_tier,
        h1_count=h1,
        h2_count=h2,
        source_format="local_text",
        augmented_with_synthetic_headers=augmented,
        injected_headers_audit=audit,
        filename=filename,
    )


def _epub_metadata_value(book, namespace: str, key: str) -> str | None:
    values = book.get_metadata(namespace, key) or []
    if not values:
        return None
    value = values[0][0] if isinstance(values[0], (tuple, list)) else values[0]
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text or None


def _parse_local_epub_document(
    raw_bytes: bytes,
    filename: str,
) -> DoclingParseResult | None:
    """Parse ordinary EPUB books locally in deterministic spine order."""

    try:
        from bs4 import BeautifulSoup
        from ebooklib import epub
    except Exception:
        return None

    temp_path = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as handle:
            handle.write(raw_bytes)
            temp_path = handle.name
        book = epub.read_epub(temp_path, options={"ignore_ncx": True})
    except Exception as exc:
        logger.warning("Local EPUB open failed for %s: %s", filename, exc)
        return None
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except OSError:
                pass

    title = _epub_metadata_value(book, "DC", "title") or Path(filename).stem
    author = _epub_metadata_value(book, "DC", "creator")
    # T-HOOK-3 — dc:language is explicit natural-language metadata; dc:date is
    # a publication-grade (but unlabelled) candidate, resolved in finalize.
    epub_language = normalize_language(_epub_metadata_value(book, "DC", "language"))
    epub_date_candidates: list[DateCandidate] = []
    dc_date = _epub_metadata_value(book, "DC", "date")
    if dc_date:
        epub_date_candidates.append(
            DateCandidate(raw=str(dc_date)[:80], kind=KIND_PUBLICATION,
                          method="epub_dc_date", source="epub:dc:date")
        )
    sections: list[Section] = []
    markdown_lines: list[str] = []
    text_blocks: list[str] = []
    h1_count = 0
    h2_count = 0
    spine_documents = 0

    for spine_entry in book.spine:
        item_id = str(spine_entry[0] if isinstance(spine_entry, (tuple, list)) else spine_entry)
        item = book.get_item_with_id(item_id)
        if item is None:
            continue
        try:
            soup = BeautifulSoup(item.get_content(), "html.parser")
        except Exception:
            continue
        for tag in soup(["script", "style", "nav", "footer", "header", "iframe"]):
            tag.decompose()

        chapter_title_tag = soup.find(["h1", "h2", "title"])
        chapter_title = re.sub(
            r"\s+", " ", chapter_title_tag.get_text(" ", strip=True)
        ).strip() if chapter_title_tag else ""
        if not chapter_title:
            chapter_title = Path(str(getattr(item, "file_name", "") or item_id)).stem
        heading_stack: list[str] = [chapter_title] if chapter_title else [title]
        chapter_has_text = False
        previous_text = ""

        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote", "pre"]):
            text = re.sub(r"\s+", " ", element.get_text(" ", strip=True)).strip()
            if not text or text == previous_text:
                continue
            previous_text = text
            if element.name and element.name.startswith("h"):
                level = max(1, min(6, int(element.name[1])))
                if level == 1:
                    h1_count += 1
                elif level == 2:
                    h2_count += 1
                heading_stack = heading_stack[: level - 1]
                while len(heading_stack) < level - 1:
                    heading_stack.append("")
                heading_stack.append(text)
                path = [part for part in heading_stack if part]
                sections.append(
                    Section(
                        heading_path=path,
                        text=text,
                        element_type="section_heading",
                        level=level,
                    )
                )
                markdown_lines.append(f"{'#' * level} {text}")
                chapter_has_text = True
                continue

            if len(text) < 2:
                continue
            path = [part for part in heading_stack if part]
            element_type = "code_block" if element.name == "pre" else "paragraph"
            sections.append(
                Section(
                    heading_path=path,
                    text=text,
                    element_type=element_type,
                )
            )
            markdown_lines.append(text)
            text_blocks.append(text)
            chapter_has_text = True

        if chapter_has_text:
            spine_documents += 1

    text = "\n\n".join(text_blocks).strip()
    if len(text) < 200 or not sections:
        logger.warning(
            "Local EPUB produced insufficient text for %s chars=%d sections=%d",
            filename,
            len(text),
            len(sections),
        )
        return None

    has_structure = (h1_count + h2_count) > 0 or spine_documents > 1
    logger.info(
        "Local EPUB parsed %s spine_docs=%d sections=%d chars=%d",
        filename,
        spine_documents,
        len(sections),
        len(text),
    )
    return DoclingParseResult(
        text=text,
        markdown="\n\n".join(markdown_lines),
        sections=sections,
        pages=None,
        has_structure=has_structure,
        source_tier=SourceTier.tier_a if has_structure else SourceTier.tier_c,
        h1_count=h1_count,
        h2_count=h2_count,
        num_pages=max(1, spine_documents),
        source_format="local_epub",
        filename=filename,
        title=title,
        author=author,
        document_date=resolve_document_dates(epub_date_candidates)["document_date"],
        language_meta=epub_language,
        date_candidates=epub_date_candidates,
        source_type="book",
    )


def _parse_local_docx_document(
    raw_bytes: bytes,
    filename: str,
) -> DoclingParseResult | None:
    """Parse DOCX headings/paragraphs locally when python-docx is installed."""
    try:
        from docx import Document
    except Exception:
        return None

    try:
        doc = Document(BytesIO(raw_bytes))
    except Exception:
        return None

    sections: list[Section] = []
    markdown_lines: list[str] = []
    text_blocks: list[str] = []
    heading_stack: list[str] = []
    h1_count = 0
    h2_count = 0

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "")
        heading_match = re.match(r"Heading\s+([1-6])$", style_name, re.IGNORECASE)
        if heading_match:
            level = int(heading_match.group(1))
            if level == 1:
                h1_count += 1
            elif level == 2:
                h2_count += 1
            heading_stack = heading_stack[: level - 1]
            while len(heading_stack) < level - 1:
                heading_stack.append("")
            heading_stack.append(text)
            path = [part for part in heading_stack if part]
            markdown_lines.append(f"{'#' * level} {text}")
            sections.append(
                Section(
                    heading_path=list(path),
                    text=text,
                    element_type="section_heading",
                    level=level,
                )
            )
            continue

        path = [part for part in heading_stack if part]
        markdown_lines.append(text)
        text_blocks.append(text)
        sections.append(
            Section(
                heading_path=list(path),
                text=text,
                element_type="paragraph",
                level=None,
            )
        )

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", (cell.text or "").strip()) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        width = max(len(row) for row in rows)
        padded = [row + [""] * (width - len(row)) for row in rows]
        header = padded[0]
        body = padded[1:]
        markdown_lines.append("| " + " | ".join(header) + " |")
        markdown_lines.append("| " + " | ".join(["---"] * width) + " |")
        for row in body:
            markdown_lines.append("| " + " | ".join(row) + " |")
        text_blocks.extend(" | ".join(row) for row in padded)
        sections.append(
            Section(
                heading_path=[part for part in heading_stack if part],
                text="\n".join(" | ".join(row) for row in padded),
                element_type="table",
                level=None,
                metadata={"table_index": table_index, "row_count": len(rows)},
            )
        )

    if not sections:
        return None

    markdown = "\n\n".join(line for line in markdown_lines if line.strip())
    text = "\n\n".join(text_blocks) or markdown
    has_structure = (h1_count + h2_count) > 0 or any(
        section.element_type == "table" for section in sections
    )
    return DoclingParseResult(
        text=text,
        markdown=markdown,
        sections=sections,
        pages=None,
        has_structure=has_structure,
        source_tier=SourceTier.tier_a if has_structure else SourceTier.tier_c,
        h1_count=h1_count,
        h2_count=h2_count,
        num_pages=1,
        source_format="local_docx",
        filename=filename,
    )


def _maybe_augment_plaintext(
    raw_bytes: bytes, filename: str, mime: str
) -> tuple[bytes, str, str, bool, list[dict]]:
    """Pre-step for plain-text uploads: try `inject_synthetic_headers` and
    re-route the (now markdown) bytes to docling under a `.md` filename so
    docling promotes the structure into `section_header` items.

    Returns:
        (bytes, filename, mime, augmented_flag, audit_list)
    """
    if not _looks_like_plain_text(filename, mime):
        return raw_bytes, filename, mime, False, []
    try:
        text = raw_bytes.decode("utf-8", errors="replace")
    except Exception:
        return raw_bytes, filename, mime, False, []

    normalized, audit = inject_synthetic_headers(text)
    if not audit:
        return raw_bytes, filename, mime, False, []

    base = filename.rsplit(".", 1)[0] if "." in filename else filename
    augmented_filename = f"{base}.md"
    audit_dicts = [
        {
            "line_no": h.line_no,
            "level": h.level,
            "pattern": h.pattern,
            "original_line": h.original_line,
        }
        for h in audit
    ]
    logger.info(
        "Pre-augmented plain text %s with %d synthetic headers → %s",
        filename, len(audit), augmented_filename,
    )
    return normalized.encode("utf-8"), augmented_filename, "text/markdown", True, audit_dicts


def _classify_tier(
    *,
    original_mime: str,
    original_filename: str,
    augmented: bool,
    h1_count: int,
    h2_count: int,
    num_pages: int,
    has_structure: bool,
) -> SourceTier:
    """Mirror the rules of the legacy source_classifier but feed off docling
    output instead of regex hits.

    Order:
      1. Multi-page PDF → ocr_ast (page-layout chunking)
      2. HTML / XHTML  → tier_b
      3. Augmented plain text with structure → tier_b_plus
      4. Native MD/DOCX with structure → tier_a
      5. Otherwise → tier_c
    """
    mime = (original_mime or "").lower()
    fname = (original_filename or "").lower()

    if mime == "application/pdf" and num_pages > 1:
        return SourceTier.ocr_ast
    if mime in ("text/html", "application/xhtml+xml") or fname.endswith((".html", ".htm", ".xhtml")):
        return SourceTier.tier_b
    if augmented and has_structure:
        return SourceTier.tier_b_plus
    if has_structure:
        return SourceTier.tier_a
    return SourceTier.tier_c


def _parse_pdf_fast_text(raw_bytes: bytes, filename: str, mime: str) -> DoclingParseResult:
    """Fast local PDF text extraction for digital PDFs when OCR is disabled.

    This preserves the adapter contract while bypassing Docling's layout/OCR
    sidecar. The chunker still gets page text, so large digital books can be
    searched quickly without paying the layout/OCR cost.
    """
    from services.ingestion.format_router import route

    decoded = route(raw_bytes, filename=filename, mime_hint=mime)
    text = decoded.text or ""
    pages = decoded.pages or ([text] if text.strip() else [])
    num_pages = max(1, len(pages))
    source_tier = SourceTier.ocr_ast if num_pages > 1 else SourceTier.tier_c
    logger.info(
        "Fast PDF text path for %s pages=%d chars=%d",
        filename,
        num_pages,
        len(text),
    )
    return DoclingParseResult(
        text=text,
        markdown=text,
        sections=[],
        pages=pages,
        has_structure=False,
        source_tier=source_tier,
        num_pages=num_pages,
        source_format="pypdf_fast_text",
        filename=filename,
    )


def _fast_pdf_text_is_usable(result: DoclingParseResult) -> bool:
    """Heuristic gate retained for diagnostics.

    Digital PDFs usually expose plenty of text via pypdf; scanned PDFs and
    image-heavy documents return empty or sparse pages. OCR is disabled, so a
    sparse result is returned as sparse text rather than falling into OCR.
    """
    text = result.text or result.markdown or ""
    compact = "".join(text.split())
    if not compact:
        return False

    pages = result.pages or [text]
    page_count = max(1, len(pages))
    nonempty_pages = [
        p for p in pages if len("".join((p or "").split())) >= _FAST_PDF_MIN_AVG_CHARS_PER_PAGE
    ]
    avg_chars = len(compact) / page_count
    nonempty_ratio = len(nonempty_pages) / page_count
    replacement_ratio = text.count("\ufffd") / max(1, len(text))
    min_total_chars = 300 if page_count <= 2 else _FAST_PDF_MIN_TOTAL_CHARS

    return (
        len(compact) >= min_total_chars
        and avg_chars >= _FAST_PDF_MIN_AVG_CHARS_PER_PAGE
        and nonempty_ratio >= _FAST_PDF_MIN_NONEMPTY_PAGE_RATIO
        and replacement_ratio <= _FAST_PDF_MAX_REPLACEMENT_RATIO
    )


def _sidecar_disabled() -> bool:
    return DOCLING_SIDECAR_POLICY in {"0", "false", "off", "disabled", "none", "local"}


def _docling_required_error(filename: str, mime: str) -> RuntimeError:
    return RuntimeError(
        "This upload needs the Docling sidecar because it is not markdown, "
        "plain text, code, HTML, or a fast-text PDF. Current parse strategy "
        f"for {filename!r} ({mime or 'unknown MIME'}) is docling_sidecar. "
        "Start it with `docker compose --profile local-parser up -d docling`, "
        "set DOCLING_SIDECAR_POLICY=auto, or convert the file to .md/.txt."
    )


async def unload_docling_sidecar() -> dict:
    """Ask the sidecar to release the heavy converter immediately."""
    async with httpx.AsyncClient(
        base_url=DOCLING_URL,
        timeout=httpx.Timeout(10.0, connect=3.0),
    ) as client:
        resp = await client.post("/unload")
        resp.raise_for_status()
        return resp.json()


async def parse_document(
    raw_bytes: bytes,
    filename: str,
    mime: str,
    do_ocr: bool = False,
) -> DoclingParseResult:
    """Hand the upload to the docling sidecar and return a structured
    DoclingParseResult. Plain-text uploads are pre-augmented with synthetic
    headers when `inject_synthetic_headers` finds qualifying markers.
    """
    if do_ocr:
        logger.warning("Ignoring do_ocr=True for %s; OCR is disabled by policy", filename)
        do_ocr = False

    # Code lane Phase 1 — early-intercept gate. Files matching a known code
    # extension (or a known filename like "Dockerfile" / "Makefile" /
    # "CMakeLists.txt") bypass the Docling sidecar entirely (no network hop,
    # no torch burn) and emit a synthetic DoclingParseResult carrying a single
    # code_block Section. tier_chunker routes this through code_splitter.pack().
    code_basename = Path(filename or "").name.lower()
    code_lang = (
        _CODE_FILENAME_TO_LANGUAGE.get(code_basename)
        or _CODE_EXT_TO_LANGUAGE.get(_extension(filename))
    )
    if code_lang:
        try:
            code_text = raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            code_text = raw_bytes.decode("utf-8", errors="replace")
        return DoclingParseResult(
            text=code_text,
            markdown=code_text,
            sections=[
                Section(
                    heading_path=[filename] if filename else [],
                    text=code_text,
                    element_type="code_block",
                    level=1,
                    language=code_lang,
                )
            ],
            pages=None,
            has_structure=False,
            source_tier=SourceTier.tier_code,
            h1_count=0,
            h2_count=0,
            num_pages=1,
            source_format=f"code_{code_lang}",
            augmented_with_synthetic_headers=False,
            injected_headers_audit=[],
            language=code_lang,
            filename=filename,
        )

    sub_result = _parse_subtitle_file(raw_bytes, filename, mime)
    if sub_result is not None:
        return sub_result

    if _looks_like_pdf(filename, mime):
        fast_result = _parse_pdf_fast_text(raw_bytes, filename, mime)
        _apply_meta(fast_result, _meta_from_pdf(raw_bytes))  # M2 pdf info
        if not do_ocr or _fast_pdf_text_is_usable(fast_result):
            return fast_result

    local_result = _parse_local_text_document(raw_bytes, filename, mime)
    if local_result is not None:
        return local_result

    aug_bytes, aug_filename, aug_mime, augmented, audit = _maybe_augment_plaintext(
        raw_bytes, filename, mime
    )

    if _sidecar_disabled():
        raise _docling_required_error(filename, mime)

    files = {"file": (aug_filename, aug_bytes, aug_mime)}
    data = {"do_ocr": "false"}

    try:
        async with httpx.AsyncClient(
            base_url=DOCLING_URL,
            timeout=httpx.Timeout(DOCLING_TIMEOUT_SECONDS, connect=30.0),
        ) as client:
            try:
                resp = await client.post("/parse", files=files, data=data)
            except httpx.RequestError as exc:
                raise RuntimeError(
                    "Docling parser sidecar is unavailable. Markdown, text, HTML, "
                    "and digital PDFs parse locally; this file type needs the "
                    "`local-parser` profile. Start it with "
                    "`docker compose --profile local-parser up -d docling`, or "
                    "convert the file to .md/.txt before ingest."
                ) from exc
            resp.raise_for_status()
            payload = resp.json()
    finally:
        if DOCLING_AUTO_UNLOAD_AFTER_PARSE:
            try:
                await unload_docling_sidecar()
            except Exception as exc:
                logger.debug("Docling sidecar auto-unload failed: %s", exc)

    sections = [
        Section(
            heading_path=list(s.get("heading_path") or []),
            text=s.get("text", "") or "",
            element_type=s.get("element_type", "paragraph"),
            level=s.get("level"),
            language=s.get("language"),
            metadata=s.get("metadata") or {},
        )
        for s in payload.get("sections", [])
    ]

    h1 = int(payload.get("h1_count", 0))
    h2 = int(payload.get("h2_count", 0))
    has_structure = bool(payload.get("has_structure", (h1 + h2) >= 2))
    num_pages = int(payload.get("num_pages", 1))

    tier = _classify_tier(
        original_mime=mime,
        original_filename=filename,
        augmented=augmented,
        h1_count=h1,
        h2_count=h2,
        num_pages=num_pages,
        has_structure=has_structure,
    )

    return DoclingParseResult(
        text=payload.get("text", "") or "",
        markdown=payload.get("markdown", "") or "",
        sections=sections,
        pages=payload.get("pages"),
        has_structure=has_structure,
        source_tier=tier,
        h1_count=h1,
        h2_count=h2,
        num_pages=num_pages,
        source_format=payload.get("source_format", "") or "",
        augmented_with_synthetic_headers=augmented,
        injected_headers_audit=audit,
        language=None,
        filename=filename,
    )
