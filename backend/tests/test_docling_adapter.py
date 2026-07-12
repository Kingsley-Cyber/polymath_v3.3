"""
Phase 7.6 — docling adapter smoke tests.

These hit the running docling sidecar over HTTP (DOCLING_URL env var). Run
inside the backend container after `docker compose up -d docling backend`:

    docker compose exec backend pytest backend/tests/test_docling_adapter.py -v

Each test exercises one classifier path:
  • Markdown with native H1/H2          → tier_a, has_structure=True
  • Plain .txt with semantic ALL CAPS   → tier_b_plus (via inject_synthetic_headers
                                          pre-augmentation), heading_path filled
  • DOCX with real heading styles       → tier_a (or tier_b_plus when augmented)
"""

from __future__ import annotations

import io
import os

import pytest

# Skip the whole module when DOCLING_URL is unset and the sidecar isn't
# reachable — keeps CI happy outside of docker compose.
DOCLING_URL = os.getenv("DOCLING_URL", "http://docling:8500")


@pytest.fixture(scope="module")
def adapter():
    from services.ingestion import docling_adapter
    return docling_adapter


def test_parser_strategy_keeps_md_txt_and_query_runtime_off_docling(adapter):
    assert adapter.parser_strategy("notes.md", "text/markdown") == "local_markdown"
    assert adapter.parser_strategy("notes.txt", "text/plain") == "local_text"
    assert adapter.parser_strategy("products.csv", "text/csv") == "local_csv"
    assert adapter.parser_strategy("inventory.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet") == "local_spreadsheet"
    assert adapter.parser_strategy("book.pdf", "application/pdf") == "local_pdf_fast_text"
    assert adapter.parser_strategy("plan.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document") == "local_docx"
    assert adapter.parser_strategy("book.epub", "application/octet-stream") == "local_epub"
    assert adapter.docling_sidecar_needed("notes.md", "text/markdown") is False
    assert adapter.docling_sidecar_needed("notes.txt", "text/plain") is False
    assert adapter.docling_sidecar_needed("products.csv", "text/csv") is False
    assert adapter.docling_sidecar_needed("inventory.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet") is False
    assert adapter.docling_sidecar_needed("plan.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document") is False
    assert adapter.docling_sidecar_needed("book.epub", "application/octet-stream") is False


def test_navigation_only_markdown_has_no_retrievable_content(adapter):
    result = adapter._parse_local_text_document(
        (
            "![](cover.jpg)\n\n"
            "[]{#nav.xhtml}\n\n"
            "[]{#chapter-1.html}\n"
        ).encode("utf-8"),
        "empty-book.md",
        "text/markdown",
    )

    assert result is not None
    assert result.sections == []
    assert adapter.retrievable_content_text(result) == ""
    assert adapter.has_retrievable_content(result) is False


def test_short_visible_markdown_remains_retrievable(adapter):
    result = adapter._parse_local_text_document(
        b"# FACS\n\nControls visible facial action units.",
        "facs.md",
        "text/markdown",
    )

    assert result is not None
    assert adapter.has_retrievable_content(result) is True


@pytest.mark.asyncio
async def test_epub_upload_parses_locally_in_spine_order(adapter, monkeypatch):
    pytest.importorskip("ebooklib")
    from ebooklib import epub

    async def fail_post(*_args, **_kwargs):
        raise AssertionError("Docling sidecar should not be called for EPUB")

    monkeypatch.setattr(adapter.httpx.AsyncClient, "post", fail_post, raising=False)
    book = epub.EpubBook()
    book.set_identifier("test-book")
    book.set_title("Consumer Behavior")
    book.set_language("en")
    book.add_author("Test Author")
    first = epub.EpubHtml(title="First", file_name="first.xhtml", lang="en")
    first.content = "<h1>Culture</h1><p>Culture shapes consumer behavior and product meaning across markets.</p>" * 4
    second = epub.EpubHtml(title="Second", file_name="second.xhtml", lang="en")
    second.content = "<h1>Commerce</h1><p>Commerce decisions depend on values, symbols, and social context.</p>" * 4
    book.add_item(first)
    book.add_item(second)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.toc = (first, second)
    book.spine = ["nav", first, second]
    buffer = io.BytesIO()
    epub.write_epub(buffer, book)

    result = await adapter.parse_document(
        raw_bytes=buffer.getvalue(),
        filename="consumer.epub",
        mime="application/octet-stream",
        do_ocr=False,
    )

    assert result.source_format == "local_epub"
    assert result.title == "Consumer Behavior"
    assert result.author == "Test Author"
    assert result.source_type == "book"
    assert result.text.index("Culture shapes") < result.text.index("Commerce decisions")
    assert len(result.sections) >= 4


@pytest.mark.asyncio
async def test_markdown_does_not_touch_docling_sidecar(adapter, monkeypatch):
    async def fail_post(*_args, **_kwargs):  # pragma: no cover - should never run
        raise AssertionError("Docling sidecar should not be called for markdown")

    monkeypatch.setattr(adapter.httpx.AsyncClient, "post", fail_post, raising=False)

    result = await adapter.parse_document(
        raw_bytes=b"# Local\n\n| A | B |\n| --- | --- |\n| x | y |\n",
        filename="local.md",
        mime="text/markdown",
        do_ocr=False,
    )

    assert result.source_format == "local_markdown"
    assert any(section.element_type == "table" for section in result.sections)


@pytest.mark.asyncio
async def test_docling_policy_off_fails_only_for_sidecar_formats(adapter, monkeypatch):
    monkeypatch.setattr(adapter, "DOCLING_SIDECAR_POLICY", "off")

    with pytest.raises(RuntimeError, match="needs the Docling sidecar"):
        await adapter.parse_document(
            raw_bytes=b"fake docx",
            filename="plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            do_ocr=False,
        )


@pytest.mark.asyncio
async def test_pdf_with_ocr_disabled_uses_fast_text_path(adapter, monkeypatch):
    """PDF + do_ocr=False bypasses Docling and uses local pypdf extraction."""
    from services.ingestion import format_router

    called = {}

    def fake_route(data: bytes, filename: str = "", mime_hint: str = ""):
        called["args"] = (data, filename, mime_hint)
        return format_router.DecodeResult(
            text="Page one text.\n\nPage two text.",
            source_mime="application/pdf",
            doc_id="fake",
            pages=["Page one text.", "Page two text."],
        )

    monkeypatch.setattr(format_router, "route", fake_route)

    result = await adapter.parse_document(
        raw_bytes=b"%PDF fake",
        filename="book.pdf",
        mime="application/pdf",
        do_ocr=False,
    )

    assert called["args"] == (b"%PDF fake", "book.pdf", "application/pdf")
    assert result.source_format == "pypdf_fast_text"
    assert result.source_tier.value == "ocr_ast"
    assert result.sections == []
    assert result.pages == ["Page one text.", "Page two text."]


@pytest.mark.asyncio
async def test_pdf_do_ocr_true_is_ignored_by_policy(adapter, monkeypatch):
    """OCR is dead-off: even legacy true values stay on the pypdf text path."""
    from services.ingestion import format_router

    def fake_route(data: bytes, filename: str = "", mime_hint: str = ""):
        return format_router.DecodeResult(
            text="Sparse but available text.",
            source_mime="application/pdf",
            doc_id="fake",
            pages=["Sparse but available text."],
        )

    async def fail_post(*_args, **_kwargs):  # pragma: no cover - should never run
        raise AssertionError("Docling sidecar should not be called for PDF OCR")

    monkeypatch.setattr(format_router, "route", fake_route)
    monkeypatch.setattr(adapter.httpx.AsyncClient, "post", fail_post, raising=False)

    result = await adapter.parse_document(
        raw_bytes=b"%PDF fake",
        filename="scanned.pdf",
        mime="application/pdf",
        do_ocr=True,
    )

    assert result.source_format == "pypdf_fast_text"
    assert result.text == "Sparse but available text."


def test_fast_pdf_text_gate_accepts_digital_pdf(adapter):
    text = " ".join(["usable digital pdf text"] * 100)
    result = adapter.DoclingParseResult(
        text=text,
        markdown=text,
        sections=[],
        pages=[text[:900], text[900:]],
        has_structure=False,
        source_tier=adapter.SourceTier.ocr_ast,
        num_pages=2,
        source_format="pypdf_fast_text",
    )

    assert adapter._fast_pdf_text_is_usable(result)


def test_fast_pdf_text_gate_rejects_sparse_scanned_pdf(adapter):
    result = adapter.DoclingParseResult(
        text="  page 1  ",
        markdown="  page 1  ",
        sections=[],
        pages=["", "page 1", ""],
        has_structure=False,
        source_tier=adapter.SourceTier.ocr_ast,
        num_pages=3,
        source_format="pypdf_fast_text",
    )

    assert not adapter._fast_pdf_text_is_usable(result)


@pytest.mark.asyncio
async def test_markdown_with_headings_is_tier_a(adapter):
    """Native MD headings → tier_a, has_structure True, sections populated."""
    md = (
        "# Project Overview\n\n"
        "This is the intro paragraph.\n\n"
        "## Goals\n\n"
        "Goal one paragraph.\n\n"
        "## Architecture\n\n"
        "Architecture paragraph.\n\n"
        "# Conclusion\n\n"
        "Wrap-up paragraph.\n"
    )
    result = await adapter.parse_document(
        raw_bytes=md.encode("utf-8"),
        filename="overview.md",
        mime="text/markdown",
        do_ocr=False,
    )
    assert result.has_structure, "expected docling to detect >=2 headings"
    assert result.source_tier.value == "tier_a", f"got {result.source_tier}"
    # Walk the sections and confirm at least one section_heading + non-empty
    # heading_path on a paragraph.
    headings = [s for s in result.sections if s.element_type == "section_heading"]
    paragraphs = [s for s in result.sections if s.element_type != "section_heading"]
    assert len(headings) >= 2
    assert any(p.heading_path for p in paragraphs), "paragraph missing heading_path"


@pytest.mark.asyncio
async def test_plaintext_caps_lines_promote_to_tier_b_plus(adapter):
    """Plain .txt with ALL CAPS standalone "headings" should be pre-augmented
    by the adapter, parsed by docling as a markdown doc, and classified as
    tier_b_plus.
    """
    txt = (
        "Onboarding:\n\n"
        "The First Glimpse onboarding flow runs in 10 minutes.\n\n"
        "ONBOARDING & COLD START (#1-3)\n\n"
        "The core problem is that your most powerful moment is locked.\n\n"
        "ALTERNATIVELY THIS CARD IS LOCAL\n\n"
        "Local fallback paragraph here.\n"
    )
    result = await adapter.parse_document(
        raw_bytes=txt.encode("utf-8"),
        filename="Onboarding.txt",
        mime="text/plain",
        do_ocr=False,
    )
    assert result.augmented_with_synthetic_headers, "adapter should pre-augment .txt"
    assert result.source_tier.value == "tier_b_plus", f"got {result.source_tier}"
    assert result.has_structure
    assert any(s.heading_path for s in result.sections if s.element_type != "section_heading"), \
        "expected at least one paragraph with a heading_path"


@pytest.mark.asyncio
async def test_docx_with_headings_classifies_with_structure(adapter):
    """A DOCX built with proper heading styles should classify as tier_a or
    tier_b_plus and yield non-empty sections. We synthesize a tiny .docx
    inline rather than shipping a binary fixture.
    """
    pytest.importorskip("docx")  # python-docx
    from docx import Document

    doc = Document()
    doc.add_heading("Quarterly Plan", level=1)
    doc.add_paragraph("This quarter we focus on retrieval quality.")
    doc.add_heading("Workstreams", level=2)
    doc.add_paragraph("Workstream A: ingestion. Workstream B: ranking.")
    doc.add_heading("Risks", level=2)
    doc.add_paragraph("Risk one: docling first-run latency.")

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    result = await adapter.parse_document(
        raw_bytes=raw,
        filename="plan.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        do_ocr=False,
    )
    assert result.has_structure, "DOCX with >=2 headings should report has_structure"
    assert result.source_tier.value in {"tier_a", "tier_b_plus"}, \
        f"unexpected tier {result.source_tier}"
    # Heading text must round-trip into the section walk.
    titles = {s.text.strip() for s in result.sections if s.element_type == "section_heading"}
    assert "Quarterly Plan" in titles


@pytest.mark.asyncio
async def test_csv_upload_parses_as_local_table_without_docling(adapter, monkeypatch):
    async def fail_post(*_args, **_kwargs):  # pragma: no cover - should never run
        raise AssertionError("Docling sidecar should not be called for CSV")

    monkeypatch.setattr(adapter.httpx.AsyncClient, "post", fail_post, raising=False)

    raw = (
        "Product,Price,Category\n"
        "Hat,19.99,Apparel\n"
        "Box2D,0,Software\n"
    ).encode("utf-8")
    result = await adapter.parse_document(
        raw_bytes=raw,
        filename="products.csv",
        mime="text/csv",
        do_ocr=False,
    )

    assert result.source_format == "local_csv"
    assert result.source_tier.value == "tier_b"
    assert result.has_structure is True
    tables = [section for section in result.sections if section.element_type == "table"]
    assert tables
    assert tables[0].metadata["columns"] == ["Product", "Price", "Category"]
    assert tables[0].metadata["row_count"] == 2
    assert "Row 1: Product=Hat; Price=19.99; Category=Apparel" in result.text
    assert "Row 2: Product=Box2D; Price=0; Category=Software" in result.text


@pytest.mark.asyncio
async def test_xlsx_upload_parses_as_local_table_without_docling(adapter, monkeypatch):
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    async def fail_post(*_args, **_kwargs):  # pragma: no cover - should never run
        raise AssertionError("Docling sidecar should not be called for XLSX")

    monkeypatch.setattr(adapter.httpx.AsyncClient, "post", fail_post, raising=False)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Inventory"
    sheet.append(["Product", "Price", "Category"])
    sheet.append(["Hat", 19.99, "Apparel"])
    sheet.append(["Box2D", 0, "Software"])
    buf = io.BytesIO()
    workbook.save(buf)

    result = await adapter.parse_document(
        raw_bytes=buf.getvalue(),
        filename="inventory.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        do_ocr=False,
    )

    assert result.source_format == "local_xlsx"
    assert result.source_tier.value == "tier_b"
    assert result.has_structure is True
    tables = [section for section in result.sections if section.element_type == "table"]
    assert tables
    assert tables[0].metadata["sheet_name"] == "Inventory"
    assert tables[0].metadata["columns"] == ["Product", "Price", "Category"]
    assert tables[0].metadata["row_count"] == 2
    assert "Section: inventory.xlsx > Inventory" in result.text
    assert "Row 1: Product=Hat; Price=19.99; Category=Apparel" in result.text
