"""Asserting tests for M2 parse-time metadata capture + routing_trace.

Run inside the backend container:
    docker exec -i polymath_v33-backend-1 python /app/tests/test_m2_metadata.py
"""

from __future__ import annotations

import os
import sys
from io import BytesIO

_BACKEND = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if _BACKEND not in sys.path:
    sys.path.insert(0, _BACKEND)

import services.ingestion.docling_adapter as da  # noqa: E402


def test_markdown_frontmatter_extracted_then_stripped():
    md = (
        "---\n"
        "title: The Art of Testing\n"
        "author: Jane Coder\n"
        "date: 2026-01-15\n"
        "extracted: 2026-03-24\n"
        "---\n"
        "# Chapter One\n\nBody paragraph with enough words to matter.\n"
    )
    res = da._parse_local_text_document(md.encode(), "art_of_testing.md", "text/markdown")
    assert res is not None
    assert res.title == "The Art of Testing"
    assert res.author == "Jane Coder"
    assert res.document_date == "2026-01-15"
    assert "title: The Art of Testing" not in res.text  # frontmatter still stripped


def test_markdown_without_frontmatter_has_no_meta():
    res = da._parse_local_text_document(b"# Plain\n\nNo frontmatter here.", "plain.md", "")
    assert res is not None and res.title is None and res.author is None


def test_html_title_and_meta_author():
    html = (
        b"<html><head><title>  Widget   Handbook </title>"
        b'<meta name="author" content="Team Docs"></head>'
        b"<body><p>content body text</p></body></html>"
    )
    res = da._parse_local_text_document(html, "widgets.html", "text/html")
    assert res is not None
    assert res.title == "Widget Handbook"
    assert res.author == "Team Docs"


def test_docx_core_properties():
    from docx import Document

    d = Document()
    d.core_properties.title = "Quarterly Report"
    d.core_properties.author = "Alice Analyst"
    d.add_heading("Section", level=1)
    d.add_paragraph("Some body text for the section.")
    buf = BytesIO()
    d.save(buf)
    res = da._parse_local_text_document(buf.getvalue(), "report.docx", "")
    assert res is not None
    assert res.title == "Quarterly Report"
    assert res.author == "Alice Analyst"


def test_pdf_info_metadata():
    from pypdf import PdfWriter

    w = PdfWriter()
    w.add_blank_page(width=72, height=72)
    w.add_metadata({"/Title": "Spec Sheet", "/Author": "Bob Builder", "/CreationDate": "D:20260210120000"})
    buf = BytesIO()
    w.write(buf)
    meta = da._meta_from_pdf(buf.getvalue())
    assert meta["title"] == "Spec Sheet"
    assert meta["author"] == "Bob Builder"
    resolution = da.resolve_document_dates(meta["date_candidates"])
    assert resolution["document_date"] is None
    assert resolution["reason"] == "file_date_only"


def test_subtitle_gets_stem_title():
    srt = "\n".join(
        f"{i+1}\n00:00:{i:02d},000 --> 00:00:{i+1:02d},500\nBOB: line {i} spoken\n" for i in range(5)
    )
    res = da._parse_subtitle_file(srt.encode(), "team_standup.srt", "")
    assert res is not None and res.title == "team_standup"


def test_finalize_fallback_title_and_trace():
    res = da._parse_local_text_document(b"just plain text content here for the parser", "my_notes-2026.txt", "")
    assert res is not None
    da.finalize_source_meta(res, "my_notes-2026.txt")
    assert res.title == "my notes 2026"                     # cleaned stem fallback
    assert res.source_type in ("text", "document")
    tr = res.routing_trace
    assert tr["title_source"] == "filename"
    for k in ("parser", "tier", "parent_strategy", "child_strategy", "has_structure"):
        assert k in tr, f"missing trace key {k}"


def test_finalize_preserves_metadata_title():
    md = "---\ntitle: Real Title\n---\n# H\n\nbody\n"
    res = da._parse_local_text_document(md.encode(), "file_name.md", "")
    da.finalize_source_meta(res, "file_name.md")
    assert res.title == "Real Title"
    assert res.routing_trace["title_source"] == "metadata"
    assert res.source_type == "markdown"


def test_finalize_idempotent():
    res = da._parse_local_text_document(b"# H\n\nbody", "x.md", "")
    da.finalize_source_meta(res, "x.md")
    t1 = dict(res.routing_trace)
    da.finalize_source_meta(res, "x.md")
    assert res.routing_trace == t1


def _run_all():
    tests = [v for k, v in sorted(globals().items()) if k.startswith("test_") and callable(v)]
    failed = 0
    for t in tests:
        try:
            t()
            print(f"PASS {t.__name__}")
        except Exception as exc:  # noqa: BLE001
            failed += 1
            print(f"FAIL {t.__name__}: {exc!r}")
    print(f"\n{len(tests) - failed}/{len(tests)} passed")
    return failed


if __name__ == "__main__":
    sys.exit(1 if _run_all() else 0)
